import io
import streamlit as st
import pandas as pd
import docx
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile


LARGE_GROUPS = ['Lecture', 'TBL', 'Interactive Sessions', 'Panel', 'Clinical Correlations']
SMALL_GROUPS = ['Workshop', 'Lab']
DATA_COLUMNS = ['CourseName', 'EvalTitle', 'EvalName', 'QuestionType', 'QuestionText', 'EvaluateeFirst', 'EvaluateeLast']
COURSES = ['Introduction to the Profession', 
           'Scientific Foundations',
           'Scientific Inquiry', 
           'Genes and Development', 
           'Attacks and Defenses', 
           'Homeostasis', 
           'Energy and Metabolism',
           'Connection to the World',
           'Across the Lifespan',
           'Professional Responsibility',
           'Populations and Methods',
           'Anatomy',
           'Clinical Skills',
           'ILCE']
FEEDBACK_TYPE = ['Large Group', 'Small Group']


def generate_feedback_info(df: pd.DataFrame) -> tuple[str, str, int]:
    """
    Finds the course name, type of feedback and the feedback number.

    Args:
        Dataframe of feedback

    Returns:
        Tuple with course name, type of feedback and the feedback number
    
    """
    course = list(set(df['CourseName'].to_list()))
    if len(course) == 1:
        course_name = course[0]
    else:
        course_name = 'Multiple courses detected'
    eval_title = list(set(df['EvalTitle'].to_list()))[0]
    if eval_title.find('Large Group') != -1:
        feedback_type = 'Large Group'
    elif eval_title.find('Workshop/Lab') != -1:
        feedback_type = 'Small Group'
    else:
        feedback_type = 'Feedback type not detected'
    eval_name = list(set(df['EvalName'].to_list()))[0]
    eval_name.strip()
    evaluation_number = int(eval_name[-1])
    return (course_name, feedback_type, evaluation_number)

def generate_schedule_for_evaluation(schedule: pd.DataFrame, evaluation_number: int, feedback_type: str) -> pd.DataFrame:
    """
    Generate a schedule of events that were assessed for the feedback type and number

    Args:
        schedule (DataFrame): DataFrame of course schedule
        evaluation_number (int): the number of the evaluation (mid-course or qualifier)
        feedback_type (str): the type of feedback (e.g, large-group or small-group)

    Returns:
        DataFrame of events which were assessed in the feedback
    
    """
    schedule['Date'] = pd.to_datetime(schedule['Date'], format='%Y-%m-%d')
    schedule['tFrom'] = pd.to_datetime(schedule['tFrom'], format='%I:%M %p')
    schedule.sort_values(by=['Date', 'tFrom'], inplace=True)
    schedule.reset_index(inplace=True)
    assessments = schedule.loc[schedule['iLearningTypeID'].str.contains('Assessment')].index.to_list()
    assessments.insert(0,0)
    schedule_for_feedback = schedule[assessments[int(evaluation_number) - 1]:assessments[evaluation_number]]
    if feedback_type == 'Large Group':
        schedule_for_feedback = schedule_for_feedback[schedule_for_feedback['iLearningTypeID'].isin(LARGE_GROUPS)]
    elif feedback_type == 'Small Group':
        schedule_for_feedback = schedule_for_feedback[schedule_for_feedback['iLearningTypeID'].isin(SMALL_GROUPS)]
    return schedule_for_feedback

def generate_questions_from_evaluation(feedback: pd.DataFrame) -> list[str]:
    """
    Generates a list of questions in the feedback form that asks students for a rating

    Args:
        feedback (DataFrame): feedback results

    Returns:
        List[str]: list of questions from feedback results that ask for a rating

    """
    # Get the questions that ask for a rating
    ratings_df = feedback.loc[feedback['QuestionType'] == 'Radio']
    radio_questions = list(set(ratings_df.loc[ratings_df['QuestionType'] == 'Radio']['QuestionText'].to_list()))

    #Eliminate questions that ask about offensive remarks or mistreatment
    questions_to_evaluate = [q for q in radio_questions if q.find('offensive remarks') == -1 and q.find('mistreatment') == -1]
    return questions_to_evaluate

def generate_faculty_names(feedback: pd.DataFrame) -> list[tuple]:

    """
    Identifies all faculty names in the feedback results and returns a list of names

    Args:
        feedback (DataFrame): feedback results

    Returns:
        List[tuple]: list of names in format (last name, first name)
    """
    first_name = feedback['EvaluateeFirst'].to_list()
    last_name = feedback['EvaluateeLast'].to_list()
    full_name = [(last_name[index], first) for index, first in enumerate(first_name)]
    full_name = list(set(full_name))
    full_name.sort()
    return(full_name)

def convert_df_to_excel(ratings: pd.DataFrame):

    """
    Converts a DataFrame to an Excel file that can be downloaded

    Args:
        ratings (DataFrame): ratings of faculty for evaluations questions
        
    """
    output = io.BytesIO()
    ratings.to_excel(output)
    return output.getvalue()

def save_comments_to_docx(comments:list[dict], course:str, feedback_type:str):

    """
    Writes a faculty names and their comments to a docx file

    Args:
        comments(dict): list of dictionaries with Name and Comments as keys
        course(str): name of the course
        feedback_type(str): the type of feedback (large-group or small group)
    
    """
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
    doc.add_heading(course + ' - ' + feedback_type, 0)
    doc.add_heading('Comments for Faculty', 1)
    for person in comments:
        doc.add_heading(person['Name'], 2)
        for comment in person['Comments']:
            doc.add_paragraph(str(comment) + '\n')
            #doc.add_paragraph()
    #comments_file = doc.save(course + ' - '+feedback_type+' Comments.docx')
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
    

def get_sessions_for_names(names:list[tuple], schedule: pd.DataFrame, sessions: pd.DataFrame) -> list[dict]:
    
    """
    Returns a list of sessions that were taught each faculty name

    Args:
        names (list): faculty names stored in tuples
        schedule (DataFrame): course schedule
        sessions: (DataFrame): schedule shows faculty names and sesssions they taught

    Returns:
        List[dict]: keys of faculty name (str) and sessions taught (list)
    
    """
    
    faculty_sessions = []
    schedule_ids = schedule['ID'].to_list()
    schedule_ids.sort()

    for name in names:
        last_name = name[0]
        first_name = name[1]
        full_name = first_name + ' ' + last_name
        sessions_taught = sessions.loc[(sessions['cLname'] == last_name) & (sessions['cFname'] == first_name)]
        session_ids = list(set(sessions_taught['ID'].to_list()))       
        session_titles = []
        
        if len(session_ids) > 0:
            for id in session_ids:
                if id in schedule_ids:
                    title = schedule.loc[schedule['ID'] == id]['cName'].to_list()[0]
                    session_titles.append(title)
        else:
            session_titles = ['No sessions found']
        faculty_sessions.append({'name': full_name, 'sessions': session_titles})
    return faculty_sessions
    

def generate_ratings_and_comments_for_name(feedback: pd.DataFrame, name:tuple, questions_to_evaluate:list[str]) -> tuple[str, list, list[str]]:
    
    """
    Returns the name of faculty member, the ratings for that faculty member for the evaluation questions, and any comments about the faculty members

    Args:
        feedback (DataFrame): feedback results
        name (tuple): name of faculty members
        questions_to_evaluation (list): list of questions in feedback results that ask for a rating

    Return:
        tuple[str, list, list]: (name of faculty, list of ratings, list of comments)
    
    """

    ratings = []
    last_name = name[0]
    first_name = name[1]
    name = first_name + ' ' + last_name
    temp_df = feedback.loc[(feedback['EvaluateeLast'] == last_name) & (feedback['EvaluateeFirst']==first_name)]
    for question in questions_to_evaluate:
        question_df = temp_df.loc[temp_df['QuestionText'] == question]
        responses = question_df['ResponseValue'].to_list()
        non_zero_responses = [response for response in responses if response > 0.0]
        if len(non_zero_responses) > 0:
            avg = round(sum(non_zero_responses)/len(non_zero_responses), 1)
            #std = float(round(np.std(non_zero_responses), 1))
            strongly_agree_or_agree_percent = round(((non_zero_responses.count(5.0) + non_zero_responses.count(4.0))/len(non_zero_responses)), 3) * 100
            poor_or_below_average_percent = round(((non_zero_responses.count(1.0) + non_zero_responses.count(2.0))/len(non_zero_responses)), 3) * 100
        else:
            avg = None
            #std = None
            strongly_agree_or_agree_percent = None
            poor_or_below_average_percent = None
        
        ratings.append(avg)
        ratings.append(len(non_zero_responses))
        ratings.append(strongly_agree_or_agree_percent)
        ratings.append(poor_or_below_average_percent)
    
    faculty_comments = temp_df.loc[temp_df['QuestionType'] == 'Text']['ResponseText'].to_list()
    faculty_comments = [comment for comment in faculty_comments if type(comment) is str and comment.find('-----') == -1]
    if len(faculty_comments) > 0:
        comments = faculty_comments
    else:
        comments = ['No comments']
    
    return (name, ratings, comments)

def process_feedback_data(feedback: pd.DataFrame, schedule: pd.DataFrame = None, faculty: pd.DataFrame = None):

    """
    Writes ratings for faculty to a Excel file and comments to docx file

    Args:
        feedback_loc (str): location of the feedback file (csv)
        schedule_loc (str): location of the schedule file (csv)
        faculty_loc (str): location of the faculty-sessions file (csv)
    
    """

    #feedback = pd.read_csv(feedback_loc)
    #(course, feedback_type, feedback_number) = generate_feedback_info(feedback)
    questions = generate_questions_from_evaluation(feedback)   
    full_names = generate_faculty_names(feedback)
    
    # Check if locations for schedule and faculty teaching are set.
    # If the are sent generate a list of dictionaries that contain
    # the name of the faculty member and the titles of the sessions
    # they taught.
    if schedule is not None and faculty is not None:
        schedule_for_eval = generate_schedule_for_evaluation(schedule, feedback_number, feedback_type)
        sessions_taught = get_sessions_for_names(full_names, schedule_for_eval, faculty)
    else:
        sessions_taught = None

    faculty_ratings = []
    faculty_comments = []
    names_for_dataframe = []
    sessions = []
    for name in full_names:
        (faculty_name, ratings, comments) = generate_ratings_and_comments_for_name(feedback, name, questions)
        names_for_dataframe.append(faculty_name)
        faculty_ratings.append(ratings)
        faculty_comments.append({'Name': faculty_name, 'Comments': comments})
        if sessions_taught is not None:
            faculty_sessions = [[session for session in item['sessions']] for item in sessions_taught if item['name'] == faculty_name][0]
            sessions.append(', '.join(faculty_sessions))

    cols = pd.MultiIndex.from_product([questions, ['Average Rating', 'Count', 'Percent Strongly Agree or Agree', 'Percent Disagree or Strongly Disagree']])
    values_and_counts_df = pd.DataFrame(faculty_ratings, index=names_for_dataframe, columns= cols)
    
    if len(sessions) == len(values_and_counts_df):
        values_and_counts_df.insert(loc=len(values_and_counts_df.columns), column='Sessions', value = sessions)
    st.header("Download Files")
    buf = io.BytesIO()

    with zipfile.ZipFile(buf, 'x') as results:
        results.writestr('Ratings.xlsx', convert_df_to_excel(values_and_counts_df))
        results.writestr('Comments.docx', save_comments_to_docx(faculty_comments, course, feedback_type))
    
    st.download_button(
        label='Download Results',
        data=buf.getvalue(),
        file_name= course+ ' '+feedback_type+' Feedback.zip',
        mime='application/zip'
    )

if __name__ == '__main__':
    COURSES.sort()
    st.title('Parse Course Evaluations')
    course = st.selectbox('Please select the course:', COURSES)
    feedback_type = st.selectbox('Please select the type of sessions for the evaluations: ', FEEDBACK_TYPE)
    feedback_number = st. selectbox('Please select the number of feedback: ', [1, 2, 3])
    st.header('Upload feedback as csv file.')
    feedback_file = st.file_uploader('')
    if feedback_file is not None:
        try:
            feedback_df = pd.read_csv(feedback_file)
        except pd.errors.ParserError:
            st.subheader("Could not parse the file. Please check the formating of the file.")
        except Exception as e:
            st.subheader(e)
        columns = feedback_df.columns
        missing_columns = [item for item in DATA_COLUMNS if item not in columns]
        if len(missing_columns) > 0:
            st.write('The file is missing these columns: ' + ', '.join(missing_columns) + '. Please check the file.')
        eval_titles = list(set(feedback_df['EvalTitle'].to_list()))
        if len(eval_titles) > 1:
            selected_title = st.selectbox('The file contains data from more than one evaluation. Please select which evaluation you would like to process: ', eval_titles)
            data_to_process = feedback_df.loc[feedback_df['EvalTitle'] == selected_title]
            #st.dataframe(data_to_process)
            st.button('Process File', on_click=process_feedback_data, args=[data_to_process])
        else:
            st.dataframe(feedback_df)
            st.button('Process File', on_click=process_feedback_data, args=[feedback_df])

        

        
